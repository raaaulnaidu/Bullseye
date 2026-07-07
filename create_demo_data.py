"""
create_demo_data.py
-------------------
Generate synthetic demo files so faculty can run BullsEye without real student data.

Usage:
    python create_demo_data.py
"""

from pathlib import Path


DEMO_DIR = Path("demo_data")
SUBMISSIONS_DIR = DEMO_DIR / "student_submissions"


def write_docx(path: Path, title: str, paragraphs: list[str]):
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("Install python-docx first: pip install python-docx") from exc

    doc = Document()
    doc.add_heading(title, level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(path)


def main():
    DEMO_DIR.mkdir(exist_ok=True)
    SUBMISSIONS_DIR.mkdir(exist_ok=True)

    write_docx(
        DEMO_DIR / "assignment_instructions.docx",
        "Demo Assignment: Retail Review Analytics Memo",
        [
            "You are a junior business analyst for a retail manager.",
            "Analyze customer review notes and a small sales table.",
            "Submit a short memo with context, table interpretation, evidence checks, a recommendation, and an AI use note.",
            "The final answer should be concise, evidence-based, and written for a non-technical manager.",
        ],
    )

    write_docx(
        DEMO_DIR / "grading_rubric.docx",
        "Demo Rubric: 20 Points",
        [
            "Context (4 pts): Identifies role, goal, audience, and constraints.",
            "Understand table (5 pts): Correctly interprets review sentiment, theme, priority, and evidence.",
            "Evidence checks (6 pts): Performs revenue or percentage calculation and links evidence to recommendation.",
            "Memo quality (4 pts): Clear situation, interpretation, recommendation, owner, and next step.",
            "AI Use Note (1 pt): States whether AI was used, what changed, and any limitations.",
        ],
    )

    submissions = {
        "alex_chen_demo.txt": """Context:
Role: junior analyst. Goal: help the store manager decide how to respond to recent review and sales patterns. Audience: retail manager. Constraint: use only the provided review and table evidence.

Understand table:
Most complaints are negative or mixed and cluster around slow checkout and product availability. Priority is high for checkout delays because it appears repeatedly and affects customer experience.

Evidence checks:
Revenue appears to decline from 8568 to 7254, a drop of about 15.3 percent. The strongest driver seems to be lower average ticket plus operational friction in checkout.

Memo:
Situation: customer experience is weakening around checkout speed and stock availability. The data suggests operational delays are hurting satisfaction and may be linked to revenue decline. Recommendation: assign the operations lead to test an express checkout lane this week and monitor review sentiment and revenue next week.

AI Use Note:
I used AI to check wording, then changed the recommendation to match the evidence. Limitation: the dataset is small.""",
        "brianna_martinez_demo.txt": """Context:
I am helping the manager understand the reviews. The audience is the store team.

Understand table:
The reviews show some bad sentiment and some positive comments. Themes include checkout, product, and staff.

Evidence checks:
Sales went down. I think it was because customers were unhappy.

Memo:
The store should improve service and make customers happier. My recommendation is to train staff and watch reviews.

AI Use Note:
AI helped with grammar.""",
        "casey_williams_demo.txt": """I looked at the assignment.

The store has issues. Some reviews are good and bad. The manager should fix the problems.

No calculations included.
No AI note.""",
    }

    for name, text in submissions.items():
        (SUBMISSIONS_DIR / name).write_text(text, encoding="utf-8")

    print(f"Demo data created in: {DEMO_DIR}")
    print("Use these files in the Streamlit app:")
    print(f"  Instructions: {DEMO_DIR / 'assignment_instructions.docx'}")
    print(f"  Rubric      : {DEMO_DIR / 'grading_rubric.docx'}")
    print(f"  Submissions : {SUBMISSIONS_DIR}")


if __name__ == "__main__":
    main()
