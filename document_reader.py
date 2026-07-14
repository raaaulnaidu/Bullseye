"""
document_reader.py
------------------
Handles text extraction from PDF, DOCX, TXT, Tableau, CSV, and Excel files.
Supports both single-file and batch (folder) reading.
"""

import csv
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Optional


def extract_text_from_pdf(file_path: str) -> str:
    """Extract plain text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts).strip()
    except ImportError:
        raise ImportError("pdfplumber is not installed. Run: pip install pdfplumber")
    except Exception as e:
        raise RuntimeError(f"Failed to read PDF '{file_path}': {e}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Extract text from tables too
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs).strip()
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX '{file_path}': {e}")


def extract_text_from_twb(file_path: str) -> str:
    """
    Extract grading-relevant metadata from a Tableau workbook (.twb).
    Pulls: worksheet names, chart titles, captions, calculated field names,
    annotations, and parameter names from the XML structure.
    """
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
    except ET.ParseError as e:
        raise RuntimeError(f"Could not parse Tableau file '{file_path}': {e}")

    sections = ["[Tableau Workbook Analysis]\n"]

    # Worksheet names and titles
    worksheets = root.findall(".//{*}worksheet") or root.findall(".//worksheet")
    if worksheets:
        sections.append("Worksheets / Charts built:")
        for ws in worksheets:
            name = ws.get("name", "")
            if name:
                sections.append(f"  - {name}")

    # Calculated fields
    calc_fields = []
    for calc in root.iter():
        if calc.get("formula") and calc.get("name"):
            calc_fields.append(f"  - {calc.get('name')}: {calc.get('formula','')[:120]}")
    if calc_fields:
        sections.append("\nCalculated fields:")
        sections.extend(calc_fields[:20])

    # Captions and titles (any element with a caption attribute)
    captions = set()
    for el in root.iter():
        for attr in ("caption", "title", "value"):
            v = el.get(attr, "").strip()
            if v and len(v) > 3 and len(v) < 200:
                captions.add(v)
    if captions:
        sections.append("\nTitles / Labels found:")
        for c in sorted(captions)[:30]:
            sections.append(f"  - {c}")

    # Data source fields
    fields = set()
    for col in root.iter():
        if col.get("datatype") and col.get("name"):
            fields.add(col.get("name", "").lstrip("[").rstrip("]"))
    if fields:
        sections.append("\nData fields used:")
        sections.append("  " + ", ".join(sorted(fields)[:40]))

    result = "\n".join(sections)
    if len(result.strip()) < 50:
        return "[Tableau file found but no readable metadata could be extracted — grade based on PDF report]"
    return result


def extract_text_from_twbx(file_path: str) -> str:
    """Extract from a Tableau packaged workbook (.twbx) — it's a ZIP containing a .twb."""
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            twb_names = [n for n in z.namelist() if n.endswith(".twb")]
            if not twb_names:
                return "[No .twb found inside .twbx — grade based on PDF report]"
            import tempfile
            with tempfile.TemporaryDirectory() as tmp:
                twb_path = z.extract(twb_names[0], tmp)
                return extract_text_from_twb(twb_path)
    except zipfile.BadZipFile:
        return "[Could not open .twbx file — it may be corrupted]"


def extract_text_from_csv(file_path: str) -> str:
    """Convert a CSV file into readable row/column text for grading."""
    rows = []
    path = Path(file_path)
    try:
        with path.open("r", encoding="utf-8-sig", newline="", errors="replace") as f:
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
            except csv.Error:
                dialect = csv.excel
            reader = csv.reader(f, dialect)
            for row in reader:
                rows.append([cell.strip() for cell in row])
    except Exception as e:
        raise RuntimeError(f"Failed to read CSV '{file_path}': {e}")

    if not rows:
        return "[CSV file is empty]"

    max_rows = 200
    max_cols = 30
    headers = rows[0][:max_cols]
    sections = [f"[CSV Submission: {path.name}]", f"Rows: {len(rows)}", ""]

    for idx, row in enumerate(rows[1:max_rows + 1], start=1):
        cells = row[:max_cols]
        if headers and len(headers) == len(cells):
            formatted = " | ".join(
                f"{headers[i] or f'Column {i+1}'}: {cells[i]}"
                for i in range(len(cells))
                if cells[i]
            )
        else:
            formatted = " | ".join(cell for cell in cells if cell)
        if formatted:
            sections.append(f"Row {idx}: {formatted}")

    if len(rows) > max_rows + 1:
        sections.append(f"\n[Truncated after {max_rows} data rows for grading context]")
    return "\n".join(sections).strip()


def extract_text_from_xlsx(file_path: str) -> str:
    """Convert an Excel .xlsx workbook into readable sheet/row/column text for grading."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl is not installed. Run: pip install openpyxl")

    path = Path(file_path)
    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
    except Exception as e:
        raise RuntimeError(f"Failed to read Excel workbook '{file_path}': {e}")

    sections = [f"[Excel Submission: {path.name}]"]
    max_rows_per_sheet = 120
    max_cols = 30

    for ws in wb.worksheets[:10]:
        sections.append(f"\n[Sheet: {ws.title}]")
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header_row = next(rows_iter)
        except StopIteration:
            sections.append("[Empty sheet]")
            continue

        headers = [
            str(v).strip() if v is not None else f"Column {i+1}"
            for i, v in enumerate(header_row[:max_cols])
        ]
        row_count = 0
        for row_count, row in enumerate(rows_iter, start=1):
            if row_count > max_rows_per_sheet:
                sections.append(f"[Truncated after {max_rows_per_sheet} rows in this sheet]")
                break
            cells = [
                "" if v is None else str(v).strip()
                for v in row[:max_cols]
            ]
            formatted = " | ".join(
                f"{headers[i]}: {cells[i]}"
                for i in range(min(len(headers), len(cells)))
                if cells[i]
            )
            if formatted:
                sections.append(f"Row {row_count}: {formatted}")
        if row_count == 0:
            sections.append("[No data rows]")

    return "\n".join(sections).strip()


def read_document(file_path: str) -> str:
    """
    Auto-detect file type and extract text content.
    Supports: .pdf, .docx, .txt, .md, .twb, .twbx, .csv, .xlsx
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        return extract_text_from_docx(file_path)
    elif suffix == ".doc":
        raise ValueError(
            "Legacy .doc files are not supported by python-docx. "
            "Please save the file as .docx or PDF before grading."
        )
    elif suffix in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    elif suffix == ".csv":
        return extract_text_from_csv(file_path)
    elif suffix == ".xlsx":
        return extract_text_from_xlsx(file_path)
    elif suffix == ".xls":
        raise ValueError(
            "Legacy .xls files are not supported. "
            "Please save the spreadsheet as .xlsx or CSV before grading."
        )
    elif suffix == ".twb":
        return extract_text_from_twb(file_path)
    elif suffix == ".twbx":
        return extract_text_from_twbx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            f"Supported: .pdf, .docx, .txt, .md, .twb, .twbx, .csv, .xlsx"
        )


def load_student_submissions(folder_path: str) -> List[Dict]:
    """
    Load all student submissions from a folder.

    Returns a list of dicts:
        [{"name": "Alice Johnson", "file": "alice_johnson.pdf", "text": "..."}, ...]
    """
    folder = Path(folder_path)
    if not folder.is_dir():
        raise NotADirectoryError(f"Submissions folder not found: {folder_path}")

    supported_extensions = {".pdf", ".docx", ".txt", ".md", ".twb", ".twbx", ".csv", ".xlsx"}
    submissions = []

    for file_path in sorted(folder.iterdir()):
        if file_path.suffix.lower() not in supported_extensions:
            continue
        if file_path.name.startswith("."):  # skip hidden files
            continue

        student_name = _infer_student_name(file_path.stem)
        try:
            text = read_document(str(file_path))
            if not text.strip():
                print(f"  ⚠️  Warning: '{file_path.name}' appears to be empty. Skipping.")
                continue
            submissions.append(
                {
                    "name": student_name,
                    "file": file_path.name,
                    "file_path": str(file_path),
                    "text": text,
                }
            )
        except Exception as e:
            print(f"  ⚠️  Could not read '{file_path.name}': {e}")

    return submissions


def _infer_student_name(stem: str) -> str:
    """
    Turn a filename stem into a readable student name.
    e.g. 'alice_johnson_assignment1' → 'Alice Johnson'
         'john-doe'                  → 'John Doe'
    """
    # Remove common suffixes like _assignment1, _submission, _hw1
    stem = re.sub(
        r"[_\-]?(assignment|submission|hw|homework|project|lab|task)\d*.*$",
        "",
        stem,
        flags=re.IGNORECASE,
    )
    # Replace underscores/hyphens with spaces and title-case
    name = stem.replace("_", " ").replace("-", " ").strip().title()
    return name if name else stem
